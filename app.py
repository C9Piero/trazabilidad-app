import io
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# 1. Configuración de la página
st.set_page_config(
    page_title="Control Interno - Pequeños Detalles",
    page_icon="🌸",
    layout="wide"
)

# 2. Credenciales
USUARIO_CORRECTO = "admin"
PASSWORD_CORRECTO = "pequenos2026"

# 3. Inicialización del estado de sesión
if "autenticado" not in st.session_state:
    st.session_state.autenticado = False

# --- BASE DE DATOS TEMPORAL EN MEMORIA DE SESIÓN ---
if "historial_proyectos" not in st.session_state:
    st.session_state.historial_proyectos = []

# --- PANTALLA DE LOGIN ---
if not st.session_state.autenticado:
    st.markdown("<h2 style='text-align: center;'>🌸 Pequeños Detalles Handmade Perú</h2>", unsafe_allow_html=True)
    st.markdown("<h4 style='text-align: center;'>Sistema Interno de Trazabilidad e Impacto</h4>", unsafe_allow_html=True)
    st.write("---")
    
    col1, col2, col3 = st.columns([1, 2, 1])
    with col2:
        st.subheader("🔐 Iniciar Sesión")
        usuario_input = st.text_input("Usuario", placeholder="Ingresa tu usuario")
        password_input = st.text_input("Contraseña", type="password", placeholder="Ingresa tu contraseña")
        
        if st.button("Ingresar al Sistema", use_container_width=True):
            if usuario_input == USUARIO_CORRECTO and password_input == PASSWORD_CORRECTO:
                st.session_state.autenticado = True
                st.success("¡Bienvenido/a! Cargando panel...")
                st.rerun()
            else:
                st.error("⚠️ Usuario o contraseña incorrectos.")

# --- APLICACIÓN PRINCIPAL (POST-LOGIN) ---
else:
    with st.sidebar:
        st.title("Pequeños Detalles")
        st.write("👤 **Usuario:** Admin")
        st.write("---")
        
        opcion_menu = st.radio(
            "Selecciona una opción:",
            ["📊 Dashboard 2026", "➕ Nuevo Proyecto", "📁 Historial de Proyectos"]
        )
        st.write("---")
        if st.button("Cerrar Sesión"):
            st.session_state.autenticado = False
            st.rerun()

    # FUNCIÓN PARA GENERAR EL WORD (.DOCX)
    def generar_word(cliente, ruc, codigo_proy, punto_origen, tipo_proyecto, fecha_ejecucion, tot_peso_recibido, tot_prod_peso, pct_aprovechamiento, impacto_co2_neto, tot_horas_generadas):
        doc = Document()

        h1 = doc.add_heading("Reporte de Impacto del Proyecto", level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p_sub = doc.add_paragraph(f"Proyecto de transformación de textiles en desuso\nCliente: {cliente}\nEmpresa: Pequeños Detalles Handmade Perú S.A.C.\nFecha: {fecha_ejecucion}")
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_heading("1. FICHA TÉCNICA DEL PROYECTO", level=2)
        t1 = doc.add_table(rows=5, cols=2)
        t1.alignment = WD_TABLE_ALIGNMENT.CENTER
        t1.style = 'Table Grid'
        datos_ficha = [
            ("Cliente", cliente),
            ("RUC", ruc),
            ("Código de Proyecto", codigo_proy),
            ("Punto de Origen", punto_origen),
            ("Tipo de Proyecto", tipo_proyecto)
        ]
        for i, (k, v) in enumerate(datos_ficha):
            t1.rows[i].cells[0].text = k
            t1.rows[i].cells[1].text = str(v)

        doc.add_heading("2. BALANCE DE MATERIAL E IMPACTO AMBIENTAL", level=2)
        t2 = doc.add_table(rows=5, cols=3)
        t2.style = 'Table Grid'
        datos_impacto = [
            ("Categoría", "Indicador", "Valor"),
            ("Ambiental", "Textiles Recibidos", f"{tot_peso_recibido:.2f} kg"),
            ("Ambiental", "Textiles Transformados", f"{tot_prod_peso:.2f} kg"),
            ("Ambiental", "Porcentaje de Aprovechamiento", f"{pct_aprovechamiento:.1f}%"),
            ("Ambiental", "Emisiones de CO₂e Evitadas (Netas)", f"{impacto_co2_neto:.2f} kg CO₂e")
        ]
        for i, (col1_val, col2_val, col3_val) in enumerate(datos_impacto):
            t2.rows[i].cells[0].text = col1_val
            t2.rows[i].cells[1].text = col2_val
            t2.rows[i].cells[2].text = col3_val

        doc.add_heading("3. IMPACTO SOCIAL Y HORAS TRABAJADAS", level=2)
        doc.add_paragraph(f"Durante el proyecto se generaron un total de {tot_horas_generadas:.1f} horas de trabajo directo repartidas entre las áreas de operaciones, corte, logística y producción descentralizada con artesanas.")

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # MENU 1: DASHBOARD CONSOLIDADOR
    if opcion_menu == "📊 Dashboard 2026":
        st.title("📊 Balance General e Indicadores 2026")
        st.markdown("Resumen de impacto ambiental y social consolidado acumulado en tiempo real.")
        
        # Calcular acumulados reales del historial
        if len(st.session_state.historial_proyectos) > 0:
            df_hist = pd.DataFrame(st.session_state.historial_proyectos)
            co2_total = df_hist["co2_neto"].sum()
            textil_total = df_hist["peso_transformado"].sum()
            horas_totales = df_hist["horas_totales"].sum()
            prod_total = df_hist["productos_unids"].sum()
        else:
            co2_total, textil_total, horas_totales, prod_total = 0.0, 0.0, 0.0, 0

        col1, col2, col3, col4 = st.columns(4)
        col1.metric("CO₂e Evitado Total", f"{co2_total:.2f} kg", "🌱 Impacto Ambiental")
        col2.metric("Textil Procesado Total", f"{textil_total:.2f} kg", "♻️ Upcycling")
        col3.metric("Horas Sociales Totales", f"{horas_totales:.1f} hrs", "👥 Confección + Corte")
        col4.metric("Productos Obtenidos", f"{prod_total:,} unids", "📦 Producción")

    # MENU 2: NUEVO PROYECTO
    elif opcion_menu == "➕ Nuevo Proyecto":
        st.title("➕ Registro de Proyecto de Upcycling")
        st.markdown("Ingresa la información requerida por cada etapa del proceso.")
        
        # ETAPA 1
        with st.expander("📌 **ETAPA 1: Recepción y Datos del Cliente**", expanded=True):
            c1, c2, c3 = st.columns(3)
            cliente = c1.text_input("Nombre del Cliente / Razón Social", value="HILTI PERÚ S.A.")
            ruc = c2.text_input("RUC del Cliente", value="20100000001")
            codigo_proy = c3.text_input("Código del Proyecto", value="HILTI-MAR26")
            
            c4, c5, c6 = st.columns(3)
            fecha_ejecucion = c4.text_input("Fecha / Período de Ejecución", value="Marzo - Abril 2026")
            punto_origen = c5.text_input("Punto de Origen (Sede/Almacén de recojo)", value="Sede Central - Lima")
            tipo_proyecto = c6.selectbox("Tipo de Proyecto", ["Upcycling de uniformes corporativos", "Transformación textil", "Donación circular"])

        # ETAPA 2
        with st.expander("👕 **ETAPA 2: Registro de Uniformes y Prendas Recibidas**", expanded=True):
            df_prendas_default = pd.DataFrame([
                {"Ítem": 1, "Descripción de Prenda": "POLERA", "Unidades": 100, "Peso Unitario (kg)": 1.0},
                {"Ítem": 2, "Descripción de Prenda": "PANTALON DRILL", "Unidades": 200, "Peso Unitario (kg)": 1.0},
                {"Ítem": 3, "Descripción de Prenda": "POLO", "Unidades": 100, "Peso Unitario (kg)": 1.0}
            ])
            df_prendas = st.data_editor(df_prendas_default, num_rows="dynamic", key="tabla_prendas")
            df_prendas["Peso Total (kg)"] = df_prendas["Unidades"] * df_prendas["Peso Unitario (kg)"]
            tot_unidades = df_prendas["Unidades"].sum()
            tot_peso_recibido = df_prendas["Peso Total (kg)"].sum()
            st.info(f"📦 **Total Prendas:** {tot_unidades} unidades | ⚖️ **Total Peso Recibido:** {tot_peso_recibido:.2f} kg")

        # ETAPA 3
        with st.expander("🔄 **ETAPA 3: Trazabilidad de Etapas de Proceso**"):
            df_trazabilidad_default = pd.DataFrame([
                {"Etapa": "Clasificación", "Fecha": "2026-03-02", "Responsable": "Área de Logística", "Peso (kg)": 280.0, "Tipo de Registro": "Registro interno"},
                {"Etapa": "Lavado", "Fecha": "2026-03-05", "Responsable": "Lavandería", "Peso (kg)": 280.0, "Tipo de Registro": "Servicio Externo"},
                {"Etapa": "Corte", "Fecha": "2026-03-06", "Responsable": "Taller de corte", "Peso (kg)": 278.0, "Tipo de Registro": "Pesaje real"},
                {"Etapa": "Confección", "Fecha": "2026-03-09", "Responsable": "Producción descentralizada", "Peso (kg)": 250.0, "Tipo de Registro": "Entrega / Recepción"}
            ])
            df_trazabilidad = st.data_editor(df_trazabilidad_default, num_rows="dynamic", key="tabla_trazabilidad")

        # ETAPA 4
        with st.expander("🎁 **ETAPA 4: Salida de Productos Elaborados**"):
            df_productos_default = pd.DataFrame([
                {"Producto": "Cartucheras", "Cantidad (Unidades)": 2000, "Peso Unitario (kg)": 0.08},
                {"Producto": "Mochilas", "Cantidad (Unidades)": 5200, "Peso Unitario (kg)": 0.11},
                {"Producto": "Bolsos", "Cantidad (Unidades)": 2000, "Peso Unitario (kg)": 0.38}
            ])
            df_productos = st.data_editor(df_productos_default, num_rows="dynamic", key="tabla_productos")
            df_productos["Peso Total (kg)"] = df_productos["Cantidad (Unidades)"] * df_productos["Peso Unitario (kg)"]
            tot_prod_unids = df_productos["Cantidad (Unidades)"].sum()
            tot_prod_peso = df_productos["Peso Total (kg)"].sum()
            st.info(f"🎒 **Total Productos Elaborados:** {tot_prod_unids} unidades | ⚖️ **Peso Transformado:** {tot_prod_peso:.2f} kg")

        # ETAPA 5
        with st.expander("👥 **ETAPA 5: Equipo de Trabajo y Generación de Horas**"):
            col_equipo1, col_equipo2 = st.columns(2)
            with col_equipo1:
                st.markdown("**Corte y Logística**")
                df_corte_default = pd.DataFrame([
                    {"Nombre": "Maria Isabel Estrada Sandoval", "Días Trabajados": 5, "Horas/Día": 8.5},
                    {"Nombre": "Genaro Jara García", "Días Trabajados": 4, "Horas/Día": 8.5},
                    {"Nombre": "Luciana Jara Estrada", "Días Trabajados": 3, "Horas/Día": 8.5}
                ])
                df_corte = st.data_editor(df_corte_default, num_rows="dynamic", key="tabla_corte")
                df_corte["Horas Totales"] = df_corte["Días Trabajados"] * df_corte["Horas/Día"]
                tot_hrs_corte = df_corte["Horas Totales"].sum()

            with col_equipo2:
                st.markdown("**Confección Descentralizada**")
                df_conf_default = pd.DataFrame([
                    {"Artesana / Nombre": "Felicita Sandoval Vílchez", "Producto": "Mochilas", "Cantidad (Unids)": 50, "Tiempo/Unid (hrs)": 0.5},
                    {"Artesana / Nombre": "Nicolle Estrada", "Producto": "Bolsos", "Cantidad (Unids)": 40, "Tiempo/Unid (hrs)": 0.75}
                ])
                df_conf = st.data_editor(df_conf_default, num_rows="dynamic", key="tabla_confeccion")
                df_conf["Horas Totales"] = df_conf["Cantidad (Unids)"] * df_conf["Tiempo/Unid (hrs)"]
                tot_hrs_conf = df_conf["Horas Totales"].sum()

        st.write("---")
        
        # PROCESAR Y GUARDAR EN EL HISTORIAL
        if st.button("🚀 Procesar Proyecto y Guardar en Historial", type="primary", use_container_width=True):
            pct_aprovechamiento = (tot_prod_peso / tot_peso_recibido * 100) if tot_peso_recibido > 0 else 0
            co2_evitado = tot_prod_peso * 7.00676
            emisiones_proceso = tot_prod_peso * 0.61976
            impacto_co2_neto = co2_evitado - emisiones_proceso
            tot_horas_generadas = tot_hrs_corte + tot_hrs_conf

            # Guardar el objeto del proyecto en el Historial de la Sesión
            nuevo_proyecto = {
                "codigo": codigo_proy,
                "cliente": cliente,
                "fecha": fecha_ejecucion,
                "peso_recibido": tot_peso_recibido,
                "peso_transformado": tot_prod_peso,
                "aprovechamiento": pct_aprovechamiento,
                "co2_neto": impacto_co2_neto,
                "horas_totales": tot_horas_generadas,
                "productos_unids": tot_prod_unids,
                "punto_origen": punto_origen,
                "tipo_proyecto": tipo_proyecto,
                "ruc": ruc
            }
            
            # Evitar duplicados con el mismo código de proyecto
            st.session_state.historial_proyectos = [p for p in st.session_state.historial_proyectos if p["codigo"] != codigo_proy]
            st.session_state.historial_proyectos.append(nuevo_proyecto)

            st.success(f"✅ ¡Proyecto **{codigo_proy}** procesado y guardado exitosamente en el Historial!")
            
            # Métricas
            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Aprovechamiento", f"{pct_aprovechamiento:.1f}%")
            m2.metric("CO₂e Neto Evitado", f"{impacto_co2_neto:.2f} kg")
            m3.metric("Total Horas Sociales", f"{tot_horas_generadas:.1f} hrs")
            m4.metric("Productos Obtenidos", f"{tot_prod_unids} unids")

            # Botón de Descarga Inmediata
            word_buffer = generar_word(cliente, ruc, codigo_proy, punto_origen, tipo_proyecto, fecha_ejecucion, tot_peso_recibido, tot_prod_peso, pct_aprovechamiento, impacto_co2_neto, tot_horas_generadas)
            
            st.write("---")
            st.download_button(
                label="📄 Descargar Informe Técnico en Word (.docx)",
                data=word_buffer,
                file_name=f"Informe_Tecnico_{codigo_proy}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    # MENU 3: HISTORIAL DE PROYECTOS REGISTRADOS
    elif opcion_menu == "📁 Historial de Proyectos":
        st.title("📁 Historial de Proyectos Registrados")
        st.markdown("Consulta tus proyectos guardados y vuelve a descargar tus informes cuando quieras.")
        
        if len(st.session_state.historial_proyectos) == 0:
            st.warning("⚠️ Todavía no hay proyectos registrados en esta sesión. Ve a la pestaña '➕ Nuevo Proyecto' para crear uno.")
        else:
            df_hist_ver = pd.DataFrame(st.session_state.historial_proyectos)
            
            # Tabla estilizada de resumen
            st.dataframe(
                df_hist_ver[["codigo", "cliente", "fecha", "peso_recibido", "peso_transformado", "co2_neto", "horas_totales"]],
                column_config={
                    "codigo": "Código",
                    "cliente": "Cliente",
                    "fecha": "Período",
                    "peso_recibido": st.column_config.NumberColumn("Kg Recibidos", format="%.2f kg"),
                    "peso_transformado": st.column_config.NumberColumn("Kg Transformados", format="%.2f kg"),
                    "co2_neto": st.column_config.NumberColumn("CO₂e Evitado", format="%.2f kg"),
                    "horas_totales": st.column_config.NumberColumn("Horas Sociales", format="%.1f hrs")
                },
                use_container_width=True
            )

            st.write("---")
            st.subheader("📥 Re-descargar Informe de Proyecto")
            
            # Selección de proyecto del historial para descargar su Word
            codigos_disponibles = [p["codigo"] for p in st.session_state.historial_proyectos]
            proy_sel_cod = st.selectbox("Selecciona un proyecto:", codigos_disponibles)
            
            proy_sel = next(p for p in st.session_state.historial_proyectos if p["codigo"] == proy_sel_cod)
            
            word_buffer_hist = generar_word(
                proy_sel["cliente"], proy_sel["ruc"], proy_sel["codigo"], 
                proy_sel["punto_origen"], proy_sel["tipo_proyecto"], proy_sel["fecha"], 
                proy_sel["peso_recibido"], proy_sel["peso_transformado"], 
                proy_sel["aprovechamiento"], proy_sel["co2_neto"], proy_sel["horas_totales"]
            )
            
            st.download_button(
                label=f"📄 Descargar Informe Técnico de {proy_sel['codigo']}",
                data=word_buffer_hist,
                file_name=f"Informe_Tecnico_{proy_sel['codigo']}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
